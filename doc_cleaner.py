import sys
import os
import zipfile
from pathlib import Path
import docx


def find_document_files(directory: str | os.PathLike[str]):
    """Return all .docx, .txt, and .pdf files directly inside the given directory."""
    dir_path = Path(directory)
    if not dir_path.is_dir():
        raise NotADirectoryError(f"'{dir_path}' is not a directory or does not exist.")

    DOCUMENT_EXTENSIONS = {".docx", ".txt", ".pdf"}

    return [
        path for path in sorted(dir_path.iterdir())
        if path.is_file()
        and path.suffix.lower() in DOCUMENT_EXTENSIONS
        and not path.name.startswith(("~$", ".")) # Ignore temporary and hidden files
    ]


def _unlink_hyperlinks(paragraphs) -> None:
    """Replace each hyperlink in the given paragraphs with plain runs, keeping its text."""
    for paragraph in paragraphs:
        for hyperlink in list(paragraph.hyperlinks):
            hyperlink_element = hyperlink._element
            parent = hyperlink_element.getparent()
            index = list(parent).index(hyperlink_element)
            for run_element in list(hyperlink_element):
                parent.insert(index, run_element)
                index += 1
            parent.remove(hyperlink_element)


def _clear_header_footer(header_footer) -> None:
    """Blank out all run text in a header or footer part."""
    for paragraph in header_footer.paragraphs:
        for run in paragraph.runs:
            run.text = ""


def _clean_document(document: docx.Document) -> None:
    """Remove headers/footers and unlink hyperlinks (keeping their text) in place."""
    for section in document.sections:
        for header_footer in (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ):
            _clear_header_footer(header_footer)
            _unlink_hyperlinks(header_footer.paragraphs)

    _unlink_hyperlinks(document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _unlink_hyperlinks(cell.paragraphs)


def clean_docx(paths: list[Path]) -> None:
    """Remove headers/footers and unlink hyperlinks (keeping their text) for each .docx path, saving each file back to itself."""
    for path in paths:
        print(path)

        # Try-except block handles files that are not true word zip archives (e.g. corrupted or open in another program).
        try:
            document = docx.Document(path)
        except (zipfile.BadZipFile, ValueError) as error:
            raise ValueError(
                f"'{path}' could not be opened as a .docx file. "
                "It may be open in another program (e.g. Microsoft Word) or corrupted."
            ) from error
        _clean_document(document)
        document.save(path)



def main() -> None:
    if len(sys.argv) < 2:
        print("Error: Please provide a directory path.")
        sys.exit(1)

    directory = sys.argv[1]
    document_paths = find_document_files(directory)
    print(document_paths)
    docx_paths = [path for path in document_paths if path.suffix.lower() == ".docx"]

    clean_docx(docx_paths)
    print(f"Cleaned {len(docx_paths)} .docx file(s) in '{directory}'.")


if __name__ == "__main__":
    main()